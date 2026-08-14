from io import BytesIO
from pathlib import Path
from typing import Any, List, Optional
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydantic import BaseModel, Field


class DocumentGenerateRequest(BaseModel):
    matter: str = Field(..., description="User factual legal issue or query")
    rights_explanation: str = Field(..., description="Explanation of legal rights")
    applicable_laws: List[dict] = Field(default_factory=list, description="List of applicable legal provisions")
    recommended_actions: List[str] = Field(default_factory=list, description="List of recommended action items")
    citations: List[str] = Field(default_factory=list, description="Statutory citations")
    disclaimer: Optional[str] = Field(default=None, description="Legal disclaimer")
    document_type: str = Field(default="legal_notice", description="Type of legal document")
    recipient_name: Optional[str] = Field(default="To Whom It May Concern", description="Recipient name")


class LegalDocumentGenerator:
    """Generates editable Microsoft Word (.docx) legal notices and documents."""

    def generate_docx(self, req: DocumentGenerateRequest) -> BytesIO:
        doc = docx.Document()

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        title = doc.add_paragraph()
        title_run = title.add_run("FORMAL LEGAL NOTICE AND DEMAND")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = "Arial"
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        sub_run = subtitle.add_run("(Under Applicable Statutory Provisions of Indian Law)")
        sub_run.italic = True
        sub_run.font.size = Pt(11)
        sub_run.font.color.rgb = RGBColor(100, 100, 100)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        doc.add_heading("1. STATEMENT OF FACTUAL MATTER", level=2)
        p_matter = doc.add_paragraph(req.matter)
        p_matter.paragraph_format.line_spacing = 1.15

        doc.add_heading("2. LEGAL RIGHTS EXPLANATION", level=2)
        p_rights = doc.add_paragraph(req.rights_explanation)
        p_rights.paragraph_format.line_spacing = 1.15

        if req.applicable_laws:
            doc.add_heading("3. APPLICABLE STATUTORY PROVISIONS", level=2)
            for law in req.applicable_laws:
                act_name = law.get("act", "Statute")
                sec_name = law.get("section", "")
                expl = law.get("explanation", "")
                p_law = doc.add_paragraph(style="List Bullet")
                r_law = p_law.add_run(f"{act_name}, {sec_name}: ")
                r_law.bold = True
                p_law.add_run(expl)

        if req.recommended_actions:
            doc.add_heading("4. FORMAL DEMANDS AND RECOMMENDED ACTIONS", level=2)
            for action in req.recommended_actions:
                p_act = doc.add_paragraph(style="List Number")
                p_act.add_run(action)

        if req.citations:
            doc.add_heading("5. STATUTORY CITATIONS AND REFERENCES", level=2)
            for cit in req.citations:
                doc.add_paragraph(cit, style="List Bullet")

        if req.disclaimer:
            doc.add_paragraph()
            p_disc = doc.add_paragraph()
            r_disc_title = p_disc.add_run("DISCLAIMER: ")
            r_disc_title.bold = True
            r_disc_title.font.size = Pt(9)
            r_disc = p_disc.add_run(req.disclaimer)
            r_disc.italic = True
            r_disc.font.size = Pt(9)
            r_disc.font.color.rgb = RGBColor(120, 120, 120)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def save_to_file(self, req: DocumentGenerateRequest, output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        buf = self.generate_docx(req)
        with open(output_path, "wb") as f:
            f.write(buf.getvalue())
        return output_path
